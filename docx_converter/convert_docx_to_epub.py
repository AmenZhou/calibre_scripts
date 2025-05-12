#!/usr/bin/env python3
"""
auto_headings.py
Detects title paragraphs in DOCX files and converts them to Heading 1 style.
Processes all DOCX files in the current directory and converts them to EPUB.
"""

import sys
import os
import subprocess
import re
from docx import Document
from docx.shared import Pt, Twips

# ---------- heuristics -------------------------------------------------
MAX_TITLE_LEN       = 120      # ignore paragraphs longer than this
MIN_TITLE_LEN       = 3        # minimum length of a title
MIN_FONT_SIZE       = Pt(14)   # treat larger font as potential heading
MIN_SPACING_BEFORE  = Pt(12)   # minimum spacing before a heading (increased)
MIN_SPACING_AFTER   = Pt(8)    # minimum spacing after a heading (increased)
MAX_WORDS_IN_TITLE  = 8        # maximum number of words in a title (reduced)

# Words that often appear on title pages but should not be chapter headings
TITLE_PAGE_WORDS = [
    'copyright', 'published', 'author', 'edition', 'all rights reserved',
    'by', 'thomas', 'decentralize'
]

# Short words that should never be considered chapters on their own
IGNORED_SHORT_WORDS = [
    "by", "of", "the", "and", "to", "in", "for", "with", "on", "at", "from", "as"
]

# List of multi-part chapter titles that should be treated as single chapters
MULTI_PART_TITLES = [
    ("An America Reconstituted", "A Once-Free World Restored"),
    ("The Revival of Common Sense", "The Case for Township Liberation"),
    ("Nullifying Unjust Governance", "Larger Governing Bodies")
]

def is_chapter_title(p, all_paragraphs=None, para_index=None):
    """Return True if paragraph p should be a chapter heading (Heading 1)."""
    text = p.text.strip()
    text_lower = text.lower()
    
    # Basic text validations
    # Ignore empty paragraphs or those that are too long
    if not text or len(text) > MAX_TITLE_LEN:
        return False
        
    # Skip paragraphs that are likely part of title page
    if any(word in text_lower for word in TITLE_PAGE_WORDS):
        return False
        
    # If it's a single short common word, it's not a chapter title
    if text_lower in IGNORED_SHORT_WORDS:
        return False
    
    # If it has too many words, it's probably not a title
    if len(text.split()) > MAX_WORDS_IN_TITLE:
        return False
        
    # If it ends with punctuation like period, question mark, it's probably a sentence
    if text.endswith('.') or text.endswith('?') or text.endswith('!'):
        return False
        
    # If it has multiple sentences (approximated by checking for periods not at the end)
    if '.' in text[:-1]:
        return False
    
    # ---- Formatting-based detection ----
    
    # Check font size - one of the most reliable indicators
    has_large_font = False
    for run in p.runs:
        if run.font.size and run.font.size >= MIN_FONT_SIZE:
            has_large_font = True
            break
    
    # Check paragraph spacing - headings often have more space around them
    has_significant_spacing = False
    if hasattr(p, 'paragraph_format'):
        spacing_before = p.paragraph_format.space_before or Pt(0)
        spacing_after = p.paragraph_format.space_after or Pt(0)
        if spacing_before >= MIN_SPACING_BEFORE or spacing_after >= MIN_SPACING_AFTER:
            has_significant_spacing = True
    
    # Check if paragraph is bold - often used for headings
    is_bold = all(run.font.bold for run in p.runs if run.text.strip())
    
    # Check context - look at surrounding paragraphs if available
    isolated_paragraph = False
    if all_paragraphs and para_index is not None:
        # Check if paragraph has empty paragraphs or short text before/after it
        prev_empty = para_index == 0 or not all_paragraphs[para_index-1].text.strip()
        next_empty = para_index == len(all_paragraphs)-1 or not all_paragraphs[para_index+1].text.strip()
        isolated_paragraph = prev_empty or next_empty
    
    # Combine formatting signals
    formatting_score = 0
    if has_large_font:
        formatting_score += 2
    if has_significant_spacing:
        formatting_score += 1
    if is_bold:
        formatting_score += 1
    if isolated_paragraph:
        formatting_score += 1
        
    # Consider a paragraph a heading if it has at least 2 formatting indicators
    return formatting_score >= 2

def should_combine_with_next(current_text, next_text):
    """Check if current heading should be combined with the next heading."""
    current_stripped = current_text.strip()
    next_stripped = next_text.strip()
    
    # Check if this combination is in our list of known multi-part titles
    for part1, part2 in MULTI_PART_TITLES:
        if (current_stripped == part1 and next_stripped == part2):
            return True
        # Also check with some flexibility for slight variations
        if (current_stripped in part1 or part1 in current_stripped) and \
           (next_stripped in part2 or part2 in next_stripped):
            return True
            
    return False

def analyze_document(doc):
    """Analyze document to identify potential chapter headings without making changes."""
    auto_chapters = []
    potential_chapters = []
    
    # Collect all paragraph formatting information first
    paragraphs = list(doc.paragraphs)
    
    # First pass: identify all potential chapter headings
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith("Heading"):
            auto_chapters.append((i, p.text))
            continue
            
        # Check for definite chapter headings
        if is_chapter_title(p, paragraphs, i):
            auto_chapters.append((i, p.text))
        
        # For potential chapters, look for paragraphs with some formatting but not enough
        elif p.text.strip() and len(p.text) <= MAX_TITLE_LEN:
            # Check for some formatting attributes
            has_some_formatting = False
            
            # Check font size
            for run in p.runs:
                if run.font.size and run.font.size >= MIN_FONT_SIZE:
                    has_some_formatting = True
                    break
            
            # Or check for bold text
            if not has_some_formatting:
                if any(run.font.bold for run in p.runs if run.text.strip()):
                    has_some_formatting = True
            
            if has_some_formatting:
                potential_chapters.append((i, p.text))
    
    # Filter out title page elements by looking at position in document
    # Take paragraphs after a reasonable threshold or the first heading with significant spacing
    threshold_index = min(15, len(paragraphs) // 10)  # Either 15 or 10% of the document
    
    # Find the first paragraph with significant spacing and large font - likely the first real chapter
    first_real_heading_idx = None
    for i, p in enumerate(paragraphs):
        if i <= threshold_index:
            continue  # Skip the title page material
            
        has_large_font = any(run.font.size and run.font.size >= MIN_FONT_SIZE for run in p.runs)
        has_spacing = False
        if hasattr(p, 'paragraph_format'):
            spacing_before = p.paragraph_format.space_before or Pt(0)
            spacing_after = p.paragraph_format.space_after or Pt(0)
            has_spacing = spacing_before >= MIN_SPACING_BEFORE or spacing_after >= MIN_SPACING_AFTER
            
        if has_large_font and has_spacing:
            first_real_heading_idx = i
            break
    
    # If we found a good first heading, use it to filter
    if first_real_heading_idx is not None:
        real_chapters = [(idx, text) for idx, text in auto_chapters if idx >= first_real_heading_idx]
    else:
        # Otherwise use the threshold
        real_chapters = [(idx, text) for idx, text in auto_chapters if idx > threshold_index]
    
    # Second pass: detect and combine multi-part chapter titles
    combined_chapters = []
    i = 0
    while i < len(real_chapters):
        current_idx, current_text = real_chapters[i]
        
        # Check if this heading should be combined with the next
        if i + 1 < len(real_chapters):
            next_idx, next_text = real_chapters[i + 1]
            
            # Check if they should be combined and are consecutive or very close
            if should_combine_with_next(current_text, next_text) and (next_idx == current_idx + 1 or next_idx == current_idx + 2):
                # Mark this as a multi-part title with special tag
                combined_chapters.append((current_idx, f"MULTIPART:{current_text}|{next_text}"))
                i += 2  # Skip the next heading since we combined it
            else:
                combined_chapters.append((current_idx, current_text))
                i += 1
        else:
            combined_chapters.append((current_idx, current_text))
            i += 1
            
    return combined_chapters, potential_chapters

def promote_headings(doc, selected_chapters=None):
    """Mark selected paragraphs as chapter headings (Heading 1)."""
    if selected_chapters is None:
        # Use auto-detection
        count = 0
        paragraphs = list(doc.paragraphs)
        for i, p in enumerate(paragraphs):
            if p.style.name.startswith("Heading"):
                continue  # already a heading
            if is_chapter_title(p, paragraphs, i):
                p.style = doc.styles["Heading 1"]
                count += 1
                print(f"  - Marked as chapter: {p.text[:50]}{'...' if len(p.text) > 50 else ''}")
        return count
    else:
        # Use selected paragraphs
        count = 0
        
        for i, (idx, text) in enumerate(selected_chapters):
            # Handle multi-part titles specially
            if text.startswith("MULTIPART:"):
                parts = text[len("MULTIPART:"):].split("|")
                if len(parts) == 2 and idx < len(doc.paragraphs):
                    # Find the second part
                    first_part, second_part = parts
                    
                    # Format the first part as Heading 1
                    p1 = doc.paragraphs[idx]
                    if not p1.style.name.startswith("Heading"):
                        p1.style = doc.styles["Heading 1"]
                        print(f"  - Marked as multi-part chapter (part 1): {p1.text[:50]}{'...' if len(p1.text) > 50 else ''}")
                        count += 1
                    
                    # Find the second part - check the next paragraph or the one after
                    if idx + 1 < len(doc.paragraphs) and doc.paragraphs[idx + 1].text.strip() == second_part:
                        p2 = doc.paragraphs[idx + 1]
                        if not p2.style.name.startswith("Heading"):
                            # Format as a subtitle or remove heading style
                            try:
                                p2.style = doc.styles["Subtitle"]
                            except:
                                # If Subtitle style doesn't exist, create a custom style
                                p2.style = doc.styles["Normal"]
                                for run in p2.runs:
                                    run.font.italic = True
                            
                            print(f"  - Marked as continuation (part 2): {p2.text[:50]}{'...' if len(p2.text) > 50 else ''}")
                    elif idx + 2 < len(doc.paragraphs) and doc.paragraphs[idx + 2].text.strip() == second_part:
                        p2 = doc.paragraphs[idx + 2]
                        if not p2.style.name.startswith("Heading"):
                            try:
                                p2.style = doc.styles["Subtitle"]
                            except:
                                p2.style = doc.styles["Normal"]
                                for run in p2.runs:
                                    run.font.italic = True
                            
                            print(f"  - Marked as continuation (part 2): {p2.text[:50]}{'...' if len(p2.text) > 50 else ''}")
            else:
                # Regular chapter title
                if idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    if not p.style.name.startswith("Heading"):
                        p.style = doc.styles["Heading 1"]
                        count += 1
                        print(f"  - Marked as chapter: {p.text[:50]}{'...' if len(p.text) > 50 else ''}")
        
        return count

def convert_to_epub(docx_file):
    """Convert a DOCX file to EPUB using ebook-convert."""
    try:
        # Create the output filename
        epub_file = os.path.splitext(docx_file)[0] + ".epub"
        
        # Construct the ebook-convert command
        cmd = [
            "ebook-convert",
            docx_file,
            epub_file,
            "--chapter", "//*[name()='h1']",
            "--level1-toc", "//*[name()='h1']",
            "--level2-toc", "//*[name()='h2']"
        ]
        
        # Add cover if it exists
        cover_file = "cover.jpg"
        if os.path.exists(cover_file):
            cmd.extend(["--cover", cover_file])
        
        # Run the conversion
        print(f"Running command: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            print(f"[✓] Successfully converted {docx_file} to {epub_file}")
            return True
        else:
            print(f"[✗] Error converting {docx_file} to EPUB:")
            print(result.stderr)
            return False
            
    except Exception as e:
        print(f"[✗] Error during EPUB conversion: {str(e)}")
        return False

def process_file(infile):
    """Process a single DOCX file and return the number of changes made."""
    try:
        doc = Document(infile)
        print(f"\nAnalyzing document structure in {infile}...")
        
        # First analyze the document without making changes
        auto_chapters, potential_chapters = analyze_document(doc)
        
        # Display detected chapters
        if auto_chapters:
            print("\nAutomatically detected chapter headings:")
            for i, (idx, text) in enumerate(auto_chapters):
                print(f"  {i+1}. {text[:50]}{'...' if len(text) > 50 else ''}")
        else:
            print("\nNo chapter headings were automatically detected.")
            
        # Use automatically detected chapters (no user interaction)
        selected_paragraphs = auto_chapters
        
        # Apply the selected chapter markings
        changed = promote_headings(doc, selected_paragraphs)
        
        outfile = os.path.splitext(infile)[0] + "_styled.docx"
        doc.save(outfile)
        print(f"[✓] Processed {infile}")
        print(f"[✓] Converted {changed} paragraphs to Heading 1")
        print(f"[✓] Saved: {outfile}")
        
        # Convert the styled file to EPUB
        print(f"\nConverting {outfile} to EPUB...")
        if convert_to_epub(outfile):
            return True
        return False
        
    except Exception as e:
        print(f"[✗] Error processing {infile}: {str(e)}")
        return False

def main():
    # Find all DOCX files in the current directory
    docx_files = [f for f in os.listdir('.') if f.endswith('.docx') and not f.endswith('_styled.docx')]
    
    if not docx_files:
        print("No DOCX files found in the current directory.")
        return

    print(f"Found {len(docx_files)} DOCX files to process.")
    
    # Process each file
    successful = 0
    for docx_file in docx_files:
        if process_file(docx_file):
            successful += 1
    
    print(f"\nSummary: Successfully processed {successful} out of {len(docx_files)} files.")

if __name__ == "__main__":
    main()
